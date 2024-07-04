import os
from fpdf import FPDF
from docx import Document
import re
import click

def clean_text(text: str) -> str:
    # Remove non-XML-compatible characters
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\xA0-\xFF]', '', text)
    return text

def pdf_to_docx_with_aspose(input_pdf_path: str, output_docx_path: str) -> None:
    try:
        import aspose.pdf as ap

        document = ap.Document(input_pdf_path)

        save_options = ap.DocSaveOptions()
        save_options.format = ap.DocSaveOptions.DocFormat.DOC_X
        save_options.mode = ap.DocSaveOptions.RecognitionMode.FLOW
        save_options.relative_horizontal_proximity = 2.5
        save_options.recognize_bullets = True

        document.save(output_docx_path, save_options)

        print(f'PDF converted to DOCX with Aspose.PDF: {output_docx_path}')
    except ImportError:
        print("Aspose.PDF library not found.")
    except FileNotFoundError as fnf_error:
        print(f"File not found error: {fnf_error}")
    except Exception as e:
        print(f"An error occurred during PDF to DOCX conversion: {e}")

def docx_to_pdf_with_fpdf(input_docx_path: str, output_pdf_path: str, font_path: str) -> None:
    font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
    try:

        doc = Document(input_docx_path)
        content = '\n'.join([para.text for para in doc.paragraphs])


        content = clean_text(content)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)


        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)

        pdf.multi_cell(0, 10, content)

        pdf.output(output_pdf_path)

        print(f'DOCX converted to PDF: {output_pdf_path}')
    except Exception as e:
        print(f"An error occurred during DOCX to PDF conversion: {e}")

def text_conversion(input_path: str, output_path: str, output_format: str) -> None:
    try:
        input_format = os.path.splitext(input_path)[1][1:]
        content = ""

        if input_format == 'pdf':
            pdf_to_docx_with_aspose(input_path, output_path)
        elif input_format == 'docx':
            font_path = 'path/to/DejaVuSans.ttf'
            docx_to_pdf_with_fpdf(input_path, output_path, font_path)
        elif input_format == 'txt':
            with open(input_path, 'r', encoding='utf-8') as file:
                content = file.read()

        content = clean_text(content)

        if output_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
        elif output_format == 'pdf' and input_format == 'txt':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, content)
            pdf.output(output_path)
        elif output_format == 'docx':
            pass
        else:
            print("Invalid output format.")
            return

        print(f'Content converted and saved to {output_path}')
    except Exception as e:
        print(f"An error occurred: {e}")

@click.command()
@click.argument('input_path', type=click.Path(exists=True))
@click.argument('output_path', type=click.Path())
@click.argument('output_format', type=click.Choice(['txt', 'pdf', 'docx']))
def convert_text_command(input_path: str, output_path: str, output_format: str) -> None:
    text_conversion(input_path, output_path, output_format)